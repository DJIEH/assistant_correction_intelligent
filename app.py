"""
app.py - Assistant de Correction Intelligent (Version Stable Sécurisée)

- Clé API sécurisée via fichier .env
- Support DOCX, PDF, Image (OCR)
- Affichage : Commentaire → Axes → Points forts → Note
- Génération rapport Word téléchargeable
"""

# ---------------- IMPORTS ----------------
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import datetime
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import os
from dotenv import load_dotenv
import google.generativeai as genai
import re

# ---------------- CONFIGURATION API ----------------

# Charge automatiquement le fichier .env du dossier courant
load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    st.error("❌ Clé API Gemini introuvable. Vérifie ton fichier .env")
    st.stop()

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")

# ---------------- LECTURE FICHIERS ----------------

def lire_docx(file):
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        st.error(f"Erreur DOCX : {e}")
        return ""

def lire_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        texte = ""
        for page in doc:
            texte += page.get_text()
        return texte
    except Exception as e:
        st.error(f"Erreur PDF : {e}")
        return ""

def lire_image(file):
    try:
        img = Image.open(file)
        texte = pytesseract.image_to_string(img, lang="fra")
        return texte
    except Exception as e:
        st.error(f"Erreur Image : {e}")
        return ""

# ---------------- ANALYSE IA ----------------

def analyser_avec_gemini(texte_ref, texte_eleve, matiere, niveau, style):

    if "Encourageante" in style:
        consigne_style = "Sois très bienveillant et valorise les efforts."
    elif "Stricte" in style:
        consigne_style = "Sois rigoureux et exigeant."
    else:
        consigne_style = "Sois neutre et objectif."

    prompt = f"""
Tu es un expert pédagogique niveau {niveau}.
Matière : {matiere}
Style : {consigne_style}

CORRIGÉ :
{texte_ref}

COPIE ÉLÈVE :
{texte_eleve}

Répond STRICTEMENT sous ce format :

Commentaire pédagogique :
...

Axes d'amélioration :
...

Points forts :
...

Note indicative :
.../20
"""

    try:
        response = model.generate_content(prompt)
        res = response.text

        # Extraction sécurisée avec regex
        comm = re.search(r"Commentaire pédagogique\s*:\s*(.*?)\n\s*Axes d'amélioration", res, re.DOTALL)
        axes = re.search(r"Axes d'amélioration\s*:\s*(.*?)\n\s*Points forts", res, re.DOTALL)
        pf = re.search(r"Points forts\s*:\s*(.*?)\n\s*Note indicative", res, re.DOTALL)
        note = re.search(r"Note indicative\s*:\s*(\d+)", res)

        commentaire = comm.group(1).strip() if comm else "Non généré"
        axes_amelioration = axes.group(1).strip() if axes else "Non généré"
        points_forts = pf.group(1).strip() if pf else "Non généré"
        note_finale = note.group(1) if note else "10"

        return commentaire, axes_amelioration, points_forts, note_finale

    except Exception as e:
        return f"Erreur IA : {e}", "", "", "00"

# ---------------- RAPPORT WORD ----------------

def creer_fiche_word(infos):
    doc = Document()

    p = doc.add_paragraph("REPUBLIQUE DE COTE D'IVOIRE\nUnion - Discipline - Travail")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("RAPPORT DE CORRECTION PEDAGOGIQUE IA", 0)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    data_rows = [
        ("Enseignant :", infos["enseignant"]),
        ("Apprenant :", infos["eleve"]),
        ("Matière / Niveau :", f"{infos['matiere']} ({infos['niveau']})"),
        ("Note indicative :", infos["note"]),
        ("Commentaire pédagogique :", infos["commentaire"]),
        ("Axes d'amélioration :", infos["axes"])
    ]

    for i, (label, value) in enumerate(data_rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value

    doc.add_paragraph(f"\nDocument généré le : {datetime.date.today()}")
    doc.add_paragraph("------------------------------------------------------------")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(
        "Pour tout besoin en intelligence artificielle, contactez le Techno Djieh "
        "au +225 0757283553 ou au mail : bitahdhiehd@gmail.com"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------- INTERFACE STREAMLIT ----------------

st.set_page_config(page_title="Assistant de Correction Intelligent", layout="wide")
st.title("📝 Assistant de Correction Intelligent")

with st.sidebar:
    st.header("Paramètres")
    enseignant = st.text_input("Enseignant", "M. Techno Djieh")
    eleve = st.text_input("Nom de l'élève")
    matiere = st.text_input("Matière")
    niveau = st.selectbox("Niveau", ["Primaire", "Secondaire", "Universitaire"])
    style = st.radio("Style", ["Encourageante", "Standard", "Stricte"])

st.divider()

col1, col2 = st.columns(2)

with col1:
    fichier_ref = st.file_uploader("Corrigé (DOCX optionnel)", type="docx")

with col2:
    type_doc = st.radio("Type copie élève", ["DOCX", "PDF", "Image"])

    if type_doc == "DOCX":
        fichier_eleve = st.file_uploader("Copie DOCX", type="docx")
    elif type_doc == "PDF":
        fichier_eleve = st.file_uploader("Copie PDF", type="pdf")
    else:
        fichier_eleve = st.file_uploader("Image devoir", type=["jpg", "png", "jpeg"])

if st.button("Lancer la correction IA"):

    if eleve and matiere and fichier_eleve:

        texte_ref = lire_docx(fichier_ref) if fichier_ref else ""

        if type_doc == "DOCX":
            texte_eleve = lire_docx(fichier_eleve)
        elif type_doc == "PDF":
            texte_eleve = lire_pdf(fichier_eleve)
        else:
            texte_eleve = lire_image(fichier_eleve)

        with st.spinner("Analyse IA en cours..."):

            commentaire, axes, points_forts, note = analyser_avec_gemini(
                texte_ref, texte_eleve, matiere, niveau, style
            )

            st.info("### 📝 Commentaire pédagogique")
            st.write(commentaire)

            st.warning("### 🔧 Axes d'amélioration")
            st.write(axes)

            st.success("### 🌟 Points forts")
            st.write(points_forts)

            st.metric("Note indicative", f"{note}/20")

            infos = {
                "enseignant": enseignant,
                "eleve": eleve,
                "matiere": matiere,
                "niveau": niveau,
                "note": note,
                "commentaire": commentaire,
                "axes": axes
            }

            docx_file = creer_fiche_word(infos)

            st.download_button(
                "📥 Télécharger la fiche DOCX",
                docx_file,
                file_name=f"Correction_{eleve}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.error("⚠️ Remplis tous les champs et ajoute la copie de l'élève.")